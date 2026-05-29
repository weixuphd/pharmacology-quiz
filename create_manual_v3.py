#!/usr/bin/env python3
"""Generate comprehensive Pharmacology_manu.docx with all chapter details and usage guides."""
import datetime, json, os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = '/Users/weixu/Downloads/pnas'
doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15
for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

# Load chapter stats
with open(os.path.join(BASE, 'chapter_stats.json')) as f:
    chapter_stats = json.load(f)

# Chapter titles
CN_TITLES = {
    1:"绪论",2:"药物效应动力学",3:"药物代谢动力学",4:"影响药物效应的因素及合理用药原则",
    5:"神经系统药理概论",6:"胆碱受体激动药",7:"抗胆碱酯酶药和胆碱酯酶复活药",8:"胆碱受体阻断药",
    9:"肾上腺素受体激动药",10:"肾上腺素受体阻断药",11:"局部麻醉药",12:"全身麻醉药",
    13:"镇静催眠药和促觉醒药",14:"抗癫痫药与抗惊厥药",15:"治疗神经退行性变性疾病药物",
    16:"抗精神失常药",17:"阿片类镇痛药、药物依赖性与药物滥用",18:"解热镇痛抗炎药",
    19:"影响其他自身活性物质的药物",20:"免疫调节药",21:"利尿药与脱水药",22:"抗心律失常药",
    23:"抗心肌缺血药",24:"调血脂药和抗动脉粥样硬化药",25:"抗高血压药",26:"抗慢性心功能不全药",
    27:"作用于血液及造血器官的药物",28:"作用于呼吸系统的药物",29:"作用于消化系统的药物",
    30:"子宫平滑肌兴奋药和抑制药",31:"肾上腺皮质激素类药物",32:"抗糖尿病药",
    33:"甲状腺激素和抗甲状腺药",34:"性激素类药与避孕药",35:"抗骨质疏松药",
    36:"抗菌药物概论",37:"β-内酰胺类抗生素及其他作用于细胞壁的抗生素",
    38:"大环内酯类、林可霉素类及其他抗生素",39:"氨基糖苷类及多黏菌素类抗生素",
    40:"四环素类及氯霉素类抗生素",41:"人工合成抗菌药",42:"抗结核病药及抗麻风药",
    43:"抗真菌药",44:"抗病毒药",45:"抗寄生虫药",46:"抗恶性肿瘤药",
}

EN_TITLES = {
    1:"Introduction to Pharmacology",2:"Pharmacodynamics",3:"Pharmacokinetics",
    4:"Factors Affecting Drug Effects and Rational Use",5:"Nervous System Pharmacology Overview",
    6:"Cholinergic Receptor Agonists",7:"Anticholinesterases and Cholinesterase Reactivators",
    8:"Cholinergic Receptor Blockers",9:"Adrenoceptor Agonists",10:"Adrenoceptor Blockers",
    11:"Local Anesthetics",12:"General Anesthetics",13:"Sedative-Hypnotics and Wakefulness-Promoting Drugs",
    14:"Antiepileptic and Anticonvulsant Drugs",15:"Drugs for Neurodegenerative Diseases",
    16:"Drugs for Mental Disorders",17:"Opioid Analgesics, Drug Dependence and Abuse",
    18:"NSAIDs and Antirheumatic Drugs",19:"Drugs Affecting Autacoids",20:"Immunomodulators",
    21:"Diuretics and Dehydrating Agents",22:"Antiarrhythmic Drugs",23:"Antianginal Drugs",
    24:"Lipid-Lowering and Antiatherosclerotic Drugs",25:"Antihypertensive Drugs",
    26:"Drugs for Chronic Heart Failure",27:"Drugs Affecting Blood and Hematopoietic Organs",
    28:"Drugs Acting on the Respiratory System",29:"Drugs Acting on the Digestive System",
    30:"Uterine Smooth Muscle Stimulants and Relaxants",31:"Adrenocortical Hormone Drugs",
    32:"Antidiabetic Drugs",33:"Thyroid Hormones and Antithyroid Drugs",
    34:"Sex Hormones and Contraceptives",35:"Anti-Osteoporosis Drugs",
    36:"General Principles of Antibacterial Drugs",37:"Beta-Lactam Antibiotics",
    38:"Macrolides, Lincomycins, and Other Antibiotics",39:"Aminoglycosides and Polymyxins",
    40:"Tetracyclines and Chloramphenicol",41:"Synthetic Antibacterial Drugs",
    42:"Antituberculosis and Antileprosy Drugs",43:"Antifungal Drugs",44:"Antiviral Drugs",
    45:"Antiparasitic Drugs",46:"Antineoplastic Drugs",
}

PARTS = [
    ("第一篇 总论 / Part 1: General Principles", range(1, 5)),
    ("第二篇 外周神经系统药理学 / Part 2: Peripheral NS Pharmacology", range(5, 13)),
    ("第三篇 中枢神经系统药理学 / Part 3: Central NS Pharmacology", range(13, 21)),
    ("第四篇 心血管系统药理学 / Part 4: Cardiovascular Pharmacology", range(21, 30)),
    ("第五篇 内分泌系统药理学 / Part 5: Endocrine Pharmacology", range(30, 36)),
    ("第七篇 化学治疗药物 / Part 7: Chemotherapeutic Drugs", range(36, 47)),
]

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('药理学题库系统\nPharmacology Quiz Bank System')
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)
doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run('使用与维护说明书\nUser & Maintenance Manual')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x38, 0xa1, 0x69)
doc.add_paragraph()
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_p.add_run(f'版本 3.0 · {datetime.date.today().strftime("%Y-%m-%d")}\n全46章 · 5,520题（中英双语各2,760题）\n46 Chapters · 5,520 Questions').font.size = Pt(9)
doc.add_page_break()

# ============================================================
# TOC
# ============================================================
doc.add_heading('目  录 / Table of Contents', level=1)
toc = [
    '一、系统概述 / System Overview',
    '二、题库设计总览 / Question Bank Design Overview',
    '  2.1 题型说明 / Question Types',
    '  2.2 CSV数据格式 / CSV Data Format',
    '  2.3 有效期机制 / Expiry Mechanism',
    '  2.4 LocalStorage 隔离机制 / Storage Isolation',
    '三、各章节题库详情 / Chapter-by-Chapter Details',
    '四、学生使用指南 / Student User Guide',
    '  4.1 通过手机浏览器访问',
    '  4.2 通过微信访问',
    '  4.3 答题操作说明',
    '  4.4 功能按钮说明',
    '五、教师使用指南 / Teacher User Guide',
    '  5.1 分发题库给学生',
    '  5.2 局域网课堂使用',
    '  5.3 查看学生进度',
    '六、代码维护指南 / Code Maintainer Guide',
    '  6.1 文件结构总览',
    '  6.2 环境准备',
    '  6.3 修改/新增题目',
    '  6.4 重新构建HTML',
    '  6.5 设置有效期',
    '  6.6 翻译英文题库',
    '  6.7 发布到GitHub Pages',
    '  6.8 创建新章节',
    '七、快速参考命令 / Quick Reference',
    '八、在线地址 / Online URLs',
    '九、常见问题 / FAQ',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ============================================================
# 一、系统概述
# ============================================================
doc.add_heading('一、系统概述 / System Overview', level=1)
doc.add_paragraph(
    '本系统为药理学全教材在线题库训练系统，覆盖药理学教材全部46章内容。系统提供中英双语版本，'
    '每章包含三种题型（是非题、单选题、简答题[多选]），总计5,520道题目（中英文各2,760题）。'
)
doc.add_paragraph(
    '系统采用离线自包含HTML格式，将完整题库（含答案）加密混淆后嵌入单个HTML文件中，无需服务器，'
    '可直接通过微信、浏览器打开。支持到期时间控制和首次使用天数限制，适合分发给学生进行自测练习。'
)
doc.add_heading('核心特性', level=2)
features = [
    '中英双语：首页支持 中文/English 一键切换，中英文题库完全独立',
    '三种题型：是非题（True/False）、单选题（Single Choice）、简答题多选（Multiple Choice）',
    '离线可用：HTML自包含所有题库数据，一次加载后可离线使用',
    '到期控制：硬性截止日期 + 首次使用后天数限制，双重保护',
    '进度独立：每章使用独立localStorage密钥，各章进度互不干扰',
    '微信兼容：--mode simple 构建模式确保微信内置浏览器正常使用',
    'GitHub Pages托管：免费HTTPS，全球CDN加速，推送即部署',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')
doc.add_page_break()

# ============================================================
# 二、题库设计总览
# ============================================================
doc.add_heading('二、题库设计总览 / Question Bank Design Overview', level=1)

doc.add_heading('2.1 题型说明 / Question Types', level=2)
type_desc = [
    ('是非题 (True/False)', '判断题，仅有两个选项："A. 正确/True"和"B. 错误/False"。学生点击按钮后立即显示结果和解析。总计552题。'),
    ('单选题 (Single Choice)', '5个选项（A-E），仅有一个正确答案。学生点击选项后立即显示结果和解析。总计1,927题。'),
    ('简答题(多选) (Multiple Choice)', '8-10个选项（A-J），有2个或更多正确答案。学生需点击多个选项后，再按"确认提交"按钮统一判断。正确答案以任意顺序选择均可。总计281题。'),
]
for t, d in type_desc:
    p = doc.add_paragraph()
    p.add_run(f'{t}：').bold = True
    p.add_run(d)

doc.add_heading('2.2 CSV数据格式 / CSV Data Format', level=2)
doc.add_paragraph('所有题库以CSV文件（UTF-8-BOM编码）存储，列结构如下：')
doc.add_paragraph('id, type, question, option_a, option_b, option_c, option_d, option_e, option_f, option_g, option_h, option_i, option_j, correct, explanation, category')
doc.add_paragraph(
    '中文CSV命名：pharma_cn_chXX.csv（如pharma_cn_ch01.csv）\n'
    '英文CSV命名：pharma_en_chXX.csv（如pharma_en_ch01.csv）\n'
    '所有CSV文件通过.gitignore保护，不会上传到公开的GitHub仓库。'
)

doc.add_heading('2.3 有效期机制 / Expiry Mechanism', level=2)
doc.add_paragraph(
    '每个HTML文件内置三重有效期检查：\n'
    '1. 硬性截止日期（--expire）：固定YYYY-MM-DD日期，超过即拒绝加载\n'
    '2. 首次使用天数（--max-days）：从学生首次打开应用算起N天后到期\n'
    '3. 完整性哈希：检测HTML文件是否被篡改，如被篡改则拒绝加载\n'
    '此外，系统还会检测时钟回拨行为（篡改系统时间），防止绕过有效期。'
)

doc.add_heading('2.4 LocalStorage隔离机制 / Storage Isolation', level=2)
doc.add_paragraph(
    '每个章节使用独立的localStorage密钥前缀，防止各章进度互相干扰：\n'
    '• 中文版：_chXX_fu, _chXX_lu, _chXX_qz（如_ch01_fu, _ch01_qz）\n'
    '• 英文版：_en_chXX_fu, _en_chXX_lu, _en_chXX_qz（如_en_ch01_fu）\n'
    '中英文版本之间、各章节之间的答题进度完全独立。'
)
doc.add_page_break()

# ============================================================
# 三、各章节题库详情
# ============================================================
doc.add_heading('三、各章节题库详情 / Chapter-by-Chapter Details', level=1)

for part_name, ch_range in PARTS:
    doc.add_heading(part_name, level=2)
    for ch in ch_range:
        if str(ch) not in chapter_stats:
            continue
        d = chapter_stats[str(ch)]
        tf = d['types'].get('是非题', 0)
        sc = d['types'].get('单选题', 0)
        mc = d['types'].get('简答题(多选)', 0)
        cn_title = CN_TITLES.get(ch, '')
        en_title = EN_TITLES.get(ch, '')

        p = doc.add_paragraph()
        p.add_run(f'第{ch}章 {cn_title}').bold = True
        p.add_run(f'\nCh{ch}: {en_title}')
        p.add_run(f'\n题量：{d["total"]}题 | 是非题：{tf} | 单选题：{sc} | 简答题(多选)：{mc}')

        # Category breakdown
        cats_str = ' | '.join([f'{c[0]}（{c[1]}题）' for c in d['cats'][:8]])
        if len(d['cats']) > 8:
            cats_str += f' ... 及其他{len(d["cats"])-8}个分类'
        p.add_run(f'\n分类分布：{cats_str}')

        # File links
        p.add_run(f'\n中文链接：ch{ch:02d}_cn.html  英文链接：ch{ch:02d}_en.html')
        doc.add_paragraph()  # spacer

doc.add_page_break()

# ============================================================
# 四、学生使用指南
# ============================================================
doc.add_heading('四、学生使用指南 / Student User Guide', level=1)

doc.add_heading('4.1 通过浏览器访问', level=2)
doc.add_paragraph(
    '1. 打开浏览器（Safari/Chrome/Firefox均可）\n'
    '2. 访问首页：https://weixuphd.github.io/pharmacology-quiz/\n'
    '3. 点击页面上方的"中文"或"English"选择语言\n'
    '4. 点击要学习的章节卡片进入题库\n'
    '5. 也可直接访问指定章节：https://weixuphd.github.io/pharmacology-quiz/chXX_cn.html'
)

doc.add_heading('4.2 通过微信访问', level=2)
doc.add_paragraph(
    '1. 教师将chXX_cn.html文件直接发送到微信群\n'
    '2. 学生在微信中点击文件即可打开\n'
    '3. 建议添加到"收藏"或"浮窗"方便下次打开\n'
    '注意：微信内置浏览器可能存在兼容性限制，所有HTML均使用simple模式构建以确保兼容。'
)

doc.add_heading('4.3 答题操作说明', level=2)
doc.add_paragraph(
    '是非题：阅读题目 → 点击"正确"或"错误"按钮 → 立即显示对错和解析\n'
    '单选题：阅读题目 → 点击A-E选项 → 立即显示对错和解析\n'
    '简答题(多选)：阅读题目 → 勾选所有你认为正确的选项 → 点击"确认提交" → 显示对错和解析\n\n'
    '导航方式：\n'
    '• 底部 ◀ ▶ 按钮前后翻题\n'
    '• 点阵导航：顶部的数字网格可快速跳转到任意题号\n'
    '• 键盘支持（电脑端）：方向键←→翻题，数字键1-5选择选项\n'
    '• 滑动支持（手机端）：左右滑动翻题\n\n'
    '颜色标记：\n'
    '• 绿色数字 = 答对  红色数字 = 答错  蓝色数字 = 当前题目'
)

doc.add_heading('4.4 功能说明', level=2)
doc.add_paragraph(
    '分类筛选：按知识分类过滤题目（如"抗菌作用机制"）\n'
    '模式切换：\n'
    '  • 顺序 — 按编号顺序答题\n'
    '  • 随机 — 随机抽取题目\n'
    '  • 错题 — 仅显示答错的题目（需先答过题）\n'
    '  • 未答 — 仅显示尚未作答的题目\n'
    '显示答案：查看当前题目的正确答案（标记为"跳过"）\n'
    '重置：清除当前章节的全部答题进度（需确认）\n'
    '统计面板：显示已答题数、正确数、正确率'
)
doc.add_page_break()

# ============================================================
# 五、教师使用指南
# ============================================================
doc.add_heading('五、教师使用指南 / Teacher User Guide', level=1)

doc.add_heading('5.1 分发题库给学生', level=2)
doc.add_paragraph(
    '方式一（推荐）：发送首页链接\n'
    '  将 https://weixuphd.github.io/pharmacology-quiz/ 发给学生\n'
    '  学生可在首页选择中英文和任意章节\n\n'
    '方式二：发送特定章节链接\n'
    '  如 https://weixuphd.github.io/pharmacology-quiz/ch36_cn.html\n'
    '  适合课堂布置特定章节作业\n\n'
    '方式三：微信直接发送HTML文件\n'
    '  将chXX_cn.html文件保存到手机，通过微信发送给班级群\n'
    '  学生直接在微信中打开（无需安装任何App）'
)

doc.add_heading('5.2 局域网课堂使用', level=2)
doc.add_paragraph(
    '如果教室没有稳定的互联网连接，可以使用局域网模式：\n'
    '1. 教师在电脑上运行：cd /path/to/pnas && ./start.sh\n'
    '2. 脚本会显示局域网IP地址（如 http://192.168.1.100:8080）\n'
    '3. 学生连接同一WiFi，在浏览器输入该IP地址即可访问\n'
    '4. Flask服务端模式：题目逐题从服务器返回，答案不提前泄露\n'
    '5. 速率限制：每IP每小时300次请求，每分钟30次，防止批量抓取'
)

doc.add_heading('5.3 进度跟踪', level=2)
doc.add_paragraph(
    '由于所有数据存储在学生的浏览器localStorage中，教师无法直接查看学生个人进度。\n'
    '建议的跟踪方式：\n'
    '1. 请学生截图统计面板（已答/正确/正确率）\n'
    '2. 使用局域网模式（Flask），可通过服务端日志查看请求量\n'
    '3. 安排课堂测验，使用线下方式验证学习效果'
)
doc.add_page_break()

# ============================================================
# 六、代码维护指南
# ============================================================
doc.add_heading('六、代码维护指南 / Code Maintainer Guide', level=1)

doc.add_heading('6.1 文件结构总览', level=2)
files = [
    ('build_mobile.py', '核心构建脚本：从CSV生成加密的自包含HTML'),
    ('build_bilingual_index.py', '生成中英双语首页index.html'),
    ('generate_pharma_cn_chXX.py', '各章中文题库生成脚本（46个文件）'),
    ('translate_chXX_YY.py', '各批中文→英文翻译脚本（8个文件）'),
    ('pharma_cn_chXX.csv / pharma_en_chXX.csv', '中英文题库CSV（各46个，.gitignore保护）'),
    ('chXX_cn.html / chXX_en.html', '中英文自包含题库HTML（各46个，部署用）'),
    ('index.html', '中英双语首页（语言切换）'),
    ('app.py', 'Flask服务端应用（局域网模式）'),
    ('start.sh', 'Flask快速启动脚本'),
    ('.gitignore', 'Git忽略规则：保护CSV题库文件'),
]
for fname, fdesc in files:
    p = doc.add_paragraph()
    p.add_run(f'{fname}：').bold = True
    p.add_run(fdesc)

doc.add_heading('6.2 环境准备', level=2)
doc.add_paragraph(
    '需要Python 3.9+，安装依赖：pip install -r requirements.txt\n'
    '如需生成.docx说明书：pip install python-docx\n'
    'GitHub推送可能需要配置SSH key或Personal Access Token'
)

doc.add_heading('6.3 修改/新增题目', level=2)
doc.add_paragraph(
    '方法一（推荐）：修改生成脚本\n'
    '  1. 编辑 generate_pharma_cn_chXX.py\n'
    '  2. 找到要修改的题目（通过id或注释定位）\n'
    '  3. 修改question/options/explanation/correct等字段\n'
    '  4. 运行 python3 generate_pharma_cn_chXX.py 重新生成CSV\n'
    '  5. 重新构建HTML（见6.4）\n\n'
    '方法二：直接编辑CSV\n'
    '  1. 用Excel或文本编辑器打开pharma_cn_chXX.csv\n'
    '  2. 修改题目内容\n'
    '  3. 保存（注意保持UTF-8编码）\n'
    '  4. 重新构建HTML（见6.4）'
)

doc.add_heading('6.4 重新构建HTML', level=2)
doc.add_paragraph('单个章节构建命令（中文）：')
cmd_cn = ('python3 build_mobile.py --csv pharma_cn_ch36.csv --output ch36_cn.html '
       '--title "第36章 抗菌药物概论" --expire 2026-12-31 --max-days 90 --mode simple')
doc.add_paragraph(cmd_cn)
doc.add_paragraph('单个章节构建命令（英文，含英文UI）：')
cmd_en = ('python3 build_mobile.py --csv pharma_en_ch36.csv --output ch36_en.html '
       '--title "Ch36 General Principles" --expire 2026-12-31 --max-days 90 '
       '--mode simple --key _en_ch36_ --lang en')
doc.add_paragraph(cmd_en)
doc.add_paragraph(
    '参数说明：\n'
    '  --csv：输入的CSV题库文件\n'
    '  --output：输出的HTML文件名\n'
    '  --title：页面标题（中文用"第XX章 XXX"，英文用"ChXX XXX"）\n'
    '  --expire：硬性截止日期（YYYY-MM-DD）\n'
    '  --max-days：首次使用后有效天数（0=不限制）\n'
    '  --mode：simple（微信兼容）/ full（压缩，文件更小）\n'
    '  --key：localStorage密钥前缀（中文默认_chXX_，英文默认_en_chXX_）\n'
    '  --lang：UI语言，en=全英文界面（默认cn中文界面）'
)

doc.add_heading('6.5 修改有效期', level=2)
doc.add_paragraph(
    '当前有效期设置：\n'
    '  --expire 2026-12-31（2026年7月1日硬性截止）\n'
    '  --max-days 90（首次使用后90天软限制）\n\n'
    '修改有效期的步骤：\n'
    '1. 批量重建所有92个HTML文件（46中文 + 46英文）：'
)
batch_cmd = '''# 中文章节批量重建
for ch in $(seq -w 1 46); do
  python3 build_mobile.py --csv pharma_cn_ch${ch}.csv --output ch${ch}_cn.html \\
    --key "_ch${ch}_" --title "Ch${ch}" --expire 2026-12-31 --max-days 90 --mode simple
done

# 英文章节批量重建
for ch in $(seq -w 1 46); do
  python3 build_mobile.py --csv pharma_en_ch${ch}.csv --output ch${ch}_en.html \\
    --key "_en_ch${ch}_" --title "Ch${ch}" --expire 2026-12-31 --max-days 90 --mode simple --lang en
done'''
doc.add_paragraph(batch_cmd)
doc.add_paragraph(
    '2. 更新首页有效期：编辑 build_bilingual_index.py，修改页脚 Expires 日期\n'
    '3. 运行 python3 build_bilingual_index.py 重建首页\n'
    '4. git add ch*.html index.html build_bilingual_index.py\n'
    '5. git commit -m "Update expiry to YYYY-MM-DD"\n'
    '6. git push origin main 推送到GitHub\n'
    '7. 学生刷新页面后自动使用新有效期\n\n'
    '注意：修改有效期会导致加密密钥变化，学生之前的答题记录会丢失！'
)

doc.add_heading('6.6 翻译与后处理英文题库', level=2)
doc.add_paragraph(
    '翻译工作流（三步）：\n\n'
    '第一步 — Google Translate翻译：\n'
    '  python3 retranslate_chapter.py pharma_cn_chXX.csv pharma_en_chXX.csv\n'
    '  读取中文CSV，逐题翻译所有字段（题目、选项、解析、分类），写入英文CSV。\n'
    '  内置0.6秒API延迟，100题章节约需60秒。\n\n'
    '第二步 — 后处理规范化：\n'
    '  python3 postprocess_en.py pharma_en_chXX.csv pharma_en_chXX.csv\n'
    '  修复内容：\n'
    '  - 缩写标准化：Pharmacodynamics (pharmacodynamics) → Pharmacodynamics (PD)\n'
    '  - True/False选项强制标准化：Correct/Error → True/False\n'
    '  - 微克符号：ug → μg\n'
    '  - 药品名/术语大写规范化：gram-positive → Gram-positive\n'
    '  - my country → China\n\n'
    '第三步 — 构建英文HTML：\n'
    '  python3 build_mobile.py --csv pharma_en_chXX.csv --output chXX_en.html \\\n'
    '      --key _en_chXX_ --title "ChXX" --expire 2026-12-31 \\\n'
    '      --max-days 90 --mode simple --lang en\n\n'
    '批量处理所有章节参见 6.7 节。'
)

doc.add_heading('6.7 批量维护命令速查', level=2)
doc.add_paragraph(
    '以下是在项目根目录 /Users/weixu/Downloads/pnas/ 下运行的所有批量命令：'
)
batch_cmds = [
    ('批量翻译全部章节（中→英）',
     'for ch in $(seq -w 1 46); do\n'
     '  python3 retranslate_chapter.py pharma_cn_ch${ch}.csv pharma_en_ch${ch}.csv\n'
     'done'),
    ('批量后处理全部英文CSV',
     'for ch in $(seq -w 1 46); do\n'
     '  python3 postprocess_en.py pharma_en_ch${ch}.csv pharma_en_ch${ch}.csv\n'
     'done'),
    ('批量重建全部中文HTML',
     'for ch in $(seq -w 1 46); do\n'
     '  python3 build_mobile.py --csv pharma_cn_ch${ch}.csv --output ch${ch}_cn.html \\\n'
     '    --key "_ch${ch}_" --title "Ch${ch}" --expire 2026-12-31 --max-days 90 --mode simple\n'
     'done'),
    ('批量重建全部英文HTML',
     'for ch in $(seq -w 1 46); do\n'
     '  python3 build_mobile.py --csv pharma_en_ch${ch}.csv --output ch${ch}_en.html \\\n'
     '    --key "_en_ch${ch}_" --title "Ch${ch}" --expire 2026-12-31 --max-days 90 --mode simple --lang en\n'
     'done'),
    ('一键修改有效期（全部92个文件）',
     '# 将所有上述命令中的 --expire 2026-12-31 替换为新日期后运行\n'
     '# 同时更新 build_bilingual_index.py 中的 Expires 日期\n'
     '# 然后运行 python3 build_bilingual_index.py 重建首页'),
    ('重建双语首页',
     'python3 build_bilingual_index.py'),
    ('提交并推送',
     'git add ch*.html index.html\n'
     'git commit -m "描述变更"\n'
     'git push origin main'),
    ('检查GitHub部署状态',
     'curl -s -o /dev/null -w "%{http_code}" \\\n'
     '  https://weixuphd.github.io/pharmacology-quiz/'),
]
for name, cmd in batch_cmds:
    p = doc.add_paragraph()
    p.add_run(f'{name}：').bold = True
    p.add_run(f'\n  {cmd}')

doc.add_heading('6.8 发布到GitHub Pages', level=2)
doc.add_paragraph(
    '完整发布流程：\n'
    '1. 运行批量构建命令（6.7节）生成所有中英文HTML\n'
    '2. 更新首页：python3 build_bilingual_index.py\n'
    '3. 提交变更：\n'
    '   git add ch*.html index.html\n'
    '   git commit -m "描述变更内容"\n'
    '4. 推送到GitHub：\n'
    '   git push origin main\n'
    '   （如遇网络问题可重试，GitHub在中国大陆偶尔不稳定）\n'
    '5. GitHub Pages自动部署，1-2分钟后生效\n'
    '6. 验证：访问 https://weixuphd.github.io/pharmacology-quiz/ 确认更新\n\n'
    '重要提醒：\n'
    '  - CSV文件受.gitignore保护，不会被推送到GitHub\n'
    '  - HTML文件推送后，学生可能因浏览器缓存仍看到旧版\n'
    '  - iOS Safari 清除缓存：设置 > Safari > 高级 > 网站数据 > github.io > 删除'
)

doc.add_heading('6.9 创建新章节', level=2)
doc.add_paragraph(
    '如需为教材新增章节（如第47章）：\n'
    '1. 创建generate_pharma_cn_ch47.py，参考现有generate_ch36_csv.py格式\n'
    '2. 运行python3 generate_pharma_cn_ch47.py生成CSV\n'
    '3. 构建中文HTML：python3 build_mobile.py --csv pharma_cn_ch47.csv --output ch47_cn.html ...\n'
    '4. 翻译英文：创建翻译脚本或加到现有翻译批次中\n'
    '5. 构建英文HTML：python3 build_mobile.py --csv pharma_en_ch47.csv --output ch47_en.html ...\n'
    '6. 更新build_bilingual_index.py添加新章节信息\n'
    '7. 运行python3 build_bilingual_index.py重新生成首页\n'
    '8. 提交并推送到GitHub'
)
doc.add_page_break()

# ============================================================
# 七、快速参考命令
# ============================================================
doc.add_heading('七、快速参考命令 / Quick Reference', level=1)

cmds = [
    ('生成中文CSV', 'python3 generate_pharma_cn_ch36.py'),
    ('翻译英文（Google Translate）', 'python3 retranslate_chapter.py pharma_cn_chXX.csv pharma_en_chXX.csv'),
    ('后处理规范化', 'python3 postprocess_en.py pharma_en_chXX.csv pharma_en_chXX.csv'),
    ('构建中文HTML', 'python3 build_mobile.py --csv pharma_cn_ch36.csv --output ch36_cn.html --key _ch36_ --title "Ch36" --expire 2026-12-31 --max-days 90 --mode simple'),
    ('构建英文HTML（含英文UI）', 'python3 build_mobile.py --csv pharma_en_ch36.csv --output ch36_en.html --key _en_ch36_ --title "Ch36" --expire 2026-12-31 --max-days 90 --mode simple --lang en'),
    ('批量重建全部中文', 'for ch in $(seq -w 1 46); do python3 build_mobile.py ...; done （详见6.7节）'),
    ('批量重建全部英文', 'for ch in $(seq -w 1 46); do python3 build_mobile.py ... --lang en; done （详见6.7节）'),
    ('生成双语首页', 'python3 build_bilingual_index.py'),
    ('修改有效期', '修改所有build命令中的 --expire 日期，重建全部HTML + 更新首页（详见6.5节）'),
    ('提交并推送', 'git add ch*.html index.html && git commit -m "..." && git push origin main'),
    ('生成说明书', 'python3 create_manual_v3.py'),
    ('启动Flask局域网', './start.sh'),
    ('安装依赖', 'pip install python-docx deep-translator'),
]
for name, cmd in cmds:
    p = doc.add_paragraph()
    p.add_run(f'{name}：').bold = True
    p.add_run(f'\n  {cmd}')
doc.add_page_break()

# ============================================================
# 八、在线地址
# ============================================================
doc.add_heading('八、在线地址 / Online URLs', level=1)
doc.add_paragraph(
    'GitHub仓库：https://github.com/weixuphd/pharmacology-quiz\n'
    '在线首页：https://weixuphd.github.io/pharmacology-quiz/\n'
    '中文链接格式：https://weixuphd.github.io/pharmacology-quiz/chXX_cn.html\n'
    '英文链接格式：https://weixuphd.github.io/pharmacology-quiz/chXX_en.html\n'
    '（XX为章节号，如01、36等，使用两位数格式）'
)
doc.add_page_break()

# ============================================================
# 九、常见问题
# ============================================================
doc.add_heading('九、常见问题 / FAQ', level=1)

faqs = [
    ('Q: 学生在微信中打开后显示空白？',
     'A: 确保使用 --mode simple 构建HTML。simple模式不使用压缩，兼容微信内置浏览器。'),
    ('Q: 如何延长题库有效期？',
     'A: 修改build命令中的 --expire 和 --max-days 参数，重新构建所有HTML并推送。CSV文件无需修改。'),
    ('Q: 题库内容是否会被学生获取答案？',
     'A: HTML中的题库经过XOR加密和base64编码，普通学生无法直接查看。如需更高安全性，请使用Flask局域网模式。'),
    ('Q: 中英文版本进度会混淆吗？',
     'A: 不会。中文使用_chXX_前缀，英文使用_en_chXX_前缀，localStorage完全隔离。'),
    ('Q: 各章之间进度会混淆吗？',
     'A: 不会。每章有独立的localStorage key（如_ch36_与_ch37_不同），答完一章不会影响其他章。'),
    ('Q: CSV文件在Excel中打开后中文乱码？',
     'A: CSV使用UTF-8-BOM编码。Excel打开时应选择"从文本/CSV导入"并指定UTF-8编码。或者在Python环境中直接编辑生成脚本。'),
    ('Q: GitHub push失败？',
     'A: 检查Token是否有效（https://github.com/settings/tokens），确认有repo权限。如果使用代理，加上 -c http.proxy=http://127.0.0.1:9981。'),
    ('Q: GitHub Pages更新后看不到变化？',
     'A: Pages部署需要1-2分钟。可以在URL后加?v=随机数强制刷新。也可检查GitHub仓库的Actions页面查看部署状态。'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    p.add_run(q).bold = True
    doc.add_paragraph(a)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(BASE, 'Pharmacology_manu.docx')
doc.save(output_path)
print(f'Manual saved: {output_path}')
